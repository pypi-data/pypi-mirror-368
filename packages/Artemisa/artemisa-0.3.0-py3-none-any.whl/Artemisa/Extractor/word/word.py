from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

class DocxExtractor:
    def __init__(self, path_docx):
        self.path = path_docx
        self.document = Document(self.path)

    def extract_all(self):
        """Extrae todo el contenido del documento, incluyendo párrafos, tablas, headers y footers"""
        content = {
            'paragraphs': self.extract_paragraphs(),
            'tables': self.extract_tables(),
            'headers': self.extract_headers(),
            'footers': self.extract_footers(),
            'properties': self.extract_document_properties()
        }
        return content

    def extract_paragraphs(self):
        """Extrae todos los párrafos con su estilo y formato"""
        paragraphs = []
        for para in self.document.paragraphs:
            if para.text.strip():  # Ignorar párrafos vacíos
                paragraph_info = {
                    'text': para.text,
                    'style': para.style.name,
                    'alignment': para.alignment,
                    'is_bold': para.runs[0].bold if para.runs else None,
                    'is_italic': para.runs[0].italic if para.runs else None
                }
                paragraphs.append(paragraph_info)
        return paragraphs

    def extract_tables(self):
        """Extrae todas las tablas del documento"""
        tables = []
        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Extraer texto y formato de la celda
                    cell_content = {
                        'text': cell.text,
                        'paragraphs': [p.text for p in cell.paragraphs]
                    }
                    row_data.append(cell_content)
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def extract_headers(self):
        """Extrae el contenido de los encabezados"""
        headers = []
        for section in self.document.sections:
            header_content = {
                'text': section.header.paragraphs[0].text if section.header else None,
                'paragraphs': [p.text for p in section.header.paragraphs] if section.header else []
            }
            headers.append(header_content)
        return headers

    def extract_footers(self):
        """Extrae el contenido de los pies de página"""
        footers = []
        for section in self.document.sections:
            footer_content = {
                'text': section.footer.paragraphs[0].text if section.footer else None,
                'paragraphs': [p.text for p in section.footer.paragraphs] if section.footer else []
            }
            footers.append(footer_content)
        return footers

    def extract_document_properties(self):
        """Extrae las propiedades del documento"""
        core_properties = self.document.core_properties
        return {
            'author': core_properties.author,
            'created': core_properties.created,
            'modified': core_properties.modified,
            'title': core_properties.title,
            'subject': core_properties.subject,
            'keywords': core_properties.keywords,
            'language': core_properties.language
        }

    def iter_block_items(self):
        """Itera sobre todos los elementos del documento en orden"""
        doc = self.document
        if isinstance(doc, _Document):
            doc_element = doc.element.body
        else:
            raise ValueError("Document expected")
        
        for child in doc_element.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    def get_document_structure(self):
        """Obtiene la estructura del documento manteniendo el orden de los elementos"""
        structure = []
        for block in self.iter_block_items():
            if isinstance(block, Paragraph):
                if block.text.strip():
                    structure.append({
                        'type': 'paragraph',
                        'content': block.text,
                        'style': block.style.name
                    })
            elif isinstance(block, Table):
                table_data = []
                for row in block.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                structure.append({
                    'type': 'table',
                    'content': table_data
                })
        return structure