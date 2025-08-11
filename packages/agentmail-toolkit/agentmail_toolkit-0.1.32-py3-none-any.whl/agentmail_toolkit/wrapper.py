from typing import Optional, Any
from abc import ABC
from pydantic import BaseModel
from agentmail import AgentMail

import io
import filetype
import pymupdf
import docx


class Attachment(BaseModel):
    text: Optional[str] = None
    error: Optional[str] = None
    file_type: Optional[str] = None


class Wrapper(ABC):
    _client: AgentMail = None

    def __init__(self, client: Optional[AgentMail] = None):
        self._client = client or AgentMail()

    def call_method(self, method_name: str, args: dict[str, Any]) -> BaseModel:
        if hasattr(self, method_name):
            return getattr(self, method_name)(**args)
        else:
            method = self._client
            for part in method_name.split("."):
                method = getattr(method, part)

            return method(**args)

    def get_attachment(self, thread_id: str, attachment_id: str):
        it = self._client.threads.get_attachment(
            thread_id=thread_id, attachment_id=attachment_id
        )
        file_bytes = b"".join(it)

        file_kind = filetype.guess(file_bytes)
        file_type = file_kind.mime if file_kind else None

        text = ""
        if file_type == "application/pdf":
            for page in pymupdf.Document(stream=file_bytes):
                text += page.get_text() + "\n"
        elif (
            file_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            for paragraph in docx.Document(io.BytesIO(file_bytes)).paragraphs:
                text += paragraph.text + "\n"
        else:
            return Attachment(
                error=f"Unsupported file type: {file_type or 'unknown'}",
                file_type=file_type,
            )

        return Attachment(text=text, file_type=file_type)
