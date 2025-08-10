"""
Tests for Feature Step 4: .docx template support and .eml output
"""

import os
import tempfile
import unittest
from unittest.mock import patch, MagicMock
from docx import Document

from src.emailer_inao.utils.document_processor import DocumentProcessor
from src.emailer_inao.utils.email_formatter import EmailFormatter
from src.emailer_inao.core.email_sender import EmailSender


class TestFeatureStep4(unittest.TestCase):
    """Test Feature Step 4 implementation."""
    
    def setUp(self):
        """Set up test environment."""
        self.temp_dir = tempfile.mkdtemp()
        self.document_processor = DocumentProcessor()
        self.email_formatter = EmailFormatter()
    
    def tearDown(self):
        """Clean up test environment."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_document_processor_initialization(self):
        """Test that DocumentProcessor initializes correctly."""
        self.assertIsNotNone(self.document_processor)
        self.assertIsNotNone(self.document_processor.template_processor)
    
    def test_email_formatter_initialization(self):
        """Test that EmailFormatter initializes correctly."""
        self.assertIsNotNone(self.email_formatter)
    
    def test_create_plain_email(self):
        """Test creating a plain text email."""
        email = self.email_formatter.create_plain_email(
            from_email="test@example.com",
            to_email="recipient@example.com",
            subject="Test Subject",
            plain_text_body="Test message body"
        )
        
        self.assertEqual(email['From'], "test@example.com")
        self.assertEqual(email['To'], "recipient@example.com")
        self.assertEqual(email['Subject'], "Test Subject")
        
        # Check that the email contains the message (may be base64 encoded)
        email_string = email.as_string()
        self.assertIn("test@example.com", email_string)
        self.assertIn("recipient@example.com", email_string)
        self.assertIn("Test Subject", email_string)
        
        # Check for base64 encoded content or plain text
        import base64
        if "VGVzdCBtZXNzYWdlIGJvZHk=" in email_string:  # base64 encoded
            self.assertIn("VGVzdCBtZXNzYWdlIGJvZHk=", email_string)
        else:
            self.assertIn("Test message body", email_string)
    
    def test_create_rich_email(self):
        """Test creating a rich HTML email."""
        html_body = "<p>This is <strong>HTML</strong> content</p>"
        plain_body = "This is plain text content"
        
        email = self.email_formatter.create_rich_email(
            from_email="test@example.com",
            to_email="recipient@example.com",
            subject="Test HTML Subject",
            html_body=html_body,
            plain_text_body=plain_body
        )
        
        self.assertEqual(email['From'], "test@example.com")
        self.assertEqual(email['To'], "recipient@example.com")
        self.assertEqual(email['Subject'], "Test HTML Subject")
        
        email_string = email.as_string()
        # Check for HTML content (may be base64 encoded)
        self.assertTrue("HTML" in email_string or "SFRNTAo=" in email_string)  # HTML or base64
        # Check for multipart structure
        self.assertIn("multipart", email_string)
    
    def test_save_as_eml(self):
        """Test saving email as .eml file."""
        email = self.email_formatter.create_plain_email(
            from_email="test@example.com",
            to_email="recipient@example.com",
            subject="Test EML",
            plain_text_body="Test EML body"
        )
        
        eml_path = os.path.join(self.temp_dir, "test.eml")
        self.email_formatter.save_as_eml(email, eml_path)
        
        self.assertTrue(os.path.exists(eml_path))
        
        with open(eml_path, 'r', encoding='utf-8') as f:
            content = f.read()
            self.assertIn("From: test@example.com", content)
            self.assertIn("To: recipient@example.com", content)
            self.assertIn("Subject: Test EML", content)
            # Check for base64 encoded content or plain text
            self.assertTrue("Test EML body" in content or "VGVzdCBFTUwgYm9keQ==" in content)
    
    def test_docx_template_validation(self):
        """Test .docx template validation."""
        # Create a simple .docx file for testing
        doc = Document()
        doc.add_paragraph("Hello {{name}}, your email is {{email}}")
        
        docx_path = os.path.join(self.temp_dir, "test.docx")
        doc.save(docx_path)
        
        # Test validation with matching fields
        available_fields = ['name', 'email', 'other']
        result = self.document_processor.validate_docx_template(docx_path, available_fields)
        
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.errors), 0)
        
        # Test validation with missing fields
        available_fields = ['name']  # missing 'email'
        result = self.document_processor.validate_docx_template(docx_path, available_fields)
        
        self.assertFalse(result.is_valid)
        self.assertGreater(len(result.errors), 0)
    
    def test_docx_template_processing(self):
        """Test processing .docx template with recipient data."""
        # Create a simple .docx file for testing
        doc = Document()
        doc.add_paragraph("Hello {{name}}, your email is {{email}}")
        
        docx_path = os.path.join(self.temp_dir, "test.docx")
        doc.save(docx_path)
        
        # Process with recipient data
        recipient_data = {'name': 'John Doe', 'email': 'john@example.com'}
        processed_doc = self.document_processor.process_docx_template(docx_path, recipient_data)
        
        # Check that the template was processed
        text_content = self.document_processor.convert_to_plain_text(processed_doc)
        self.assertIn("Hello John Doe", text_content)
        self.assertIn("john@example.com", text_content)
        self.assertNotIn("{{name}}", text_content)
        self.assertNotIn("{{email}}", text_content)
    
    def test_docx_to_html_conversion(self):
        """Test converting .docx to HTML."""
        # Create a simple .docx file for testing
        doc = Document()
        paragraph = doc.add_paragraph("This is a test paragraph")
        run = paragraph.runs[0]
        run.bold = True
        
        # Convert to HTML
        html_content = self.document_processor.convert_to_html(doc)
        
        self.assertIn("<p", html_content)
        self.assertIn("This is a test paragraph", html_content)
        self.assertIn("</p>", html_content)
    
    @patch('src.emailer_inao.core.email_sender.EmailSender._load_recipients')
    @patch('src.emailer_inao.core.email_sender.EmailSender._load_template')
    def test_email_sender_template_detection(self, mock_load_template, mock_load_recipients):
        """Test that EmailSender can detect .docx templates."""
        # Mock recipients
        mock_load_recipients.return_value = [{'name': 'Test', 'email': 'test@example.com'}]
        mock_load_template.return_value = "Subject: {{name}}"
        
        # Create test campaign folder with .docx template
        campaign_folder = os.path.join(self.temp_dir, "campaign")
        os.makedirs(campaign_folder)
        
        # Create a .docx template
        doc = Document()
        doc.add_paragraph("Hello {{name}}")
        docx_path = os.path.join(campaign_folder, "msg.docx")
        doc.save(docx_path)
        
        # Test template detection
        email_sender = EmailSender(campaign_folder)
        template_type, template_path = email_sender._detect_message_template()
        
        self.assertEqual(template_type, 'docx')
        self.assertEqual(template_path, docx_path)
    
    def test_message_template_conflict_detection(self):
        """Test detection of conflicting message templates."""
        # Create test campaign folder with both .txt and .docx templates
        campaign_folder = os.path.join(self.temp_dir, "campaign")
        os.makedirs(campaign_folder)
        
        # Create both templates
        with open(os.path.join(campaign_folder, "msg.txt"), 'w') as f:
            f.write("Hello {{name}}")
        
        doc = Document()
        doc.add_paragraph("Hello {{name}}")
        doc.save(os.path.join(campaign_folder, "msg.docx"))
        
        # Test that conflict is detected
        email_sender = EmailSender(campaign_folder)
        
        with self.assertRaises(ValueError) as context:
            email_sender._detect_message_template()
        
        self.assertIn("Both msg.txt and msg.docx exist", str(context.exception))


if __name__ == '__main__':
    unittest.main()
