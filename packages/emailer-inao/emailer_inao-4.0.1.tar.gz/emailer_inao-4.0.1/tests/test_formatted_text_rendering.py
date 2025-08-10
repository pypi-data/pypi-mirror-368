"""
Tests for formatted text rendering in picture generation.
"""

import os
import tempfile
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from src.emailer_inao.utils.document_processor import DocumentProcessor


class TestFormattedTextRendering:
    """Test cases for formatted text rendering."""
    
    def setup_method(self):
        """Set up test environment before each test."""
        self.temp_dir = tempfile.mkdtemp()
        self.document_processor = DocumentProcessor()
    
    def teardown_method(self):
        """Clean up test environment after each test."""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
    
    def create_formatted_docx(self, filename: str = "test.docx"):
        """Create a .docx file with various formatting."""
        doc = Document()
        
        # Bold text
        bold_para = doc.add_paragraph()
        bold_run = bold_para.add_run('Bold {{name}}')
        bold_run.bold = True
        bold_run.font.size = Pt(14)
        
        # Italic text
        italic_para = doc.add_paragraph()
        italic_run = italic_para.add_run('Italic {{company}}')
        italic_run.italic = True
        italic_run.font.size = Pt(12)
        
        # Colored text
        color_para = doc.add_paragraph()
        color_run = color_para.add_run('Red {{email}}')
        color_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Centered text
        center_para = doc.add_paragraph()
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_run = center_para.add_run('Centered Text')
        center_run.font.size = Pt(16)
        
        # Mixed formatting
        mixed_para = doc.add_paragraph()
        normal_run = mixed_para.add_run('Normal ')
        bold_run = mixed_para.add_run('Bold ')
        bold_run.bold = True
        italic_run = mixed_para.add_run('Italic')
        italic_run.italic = True
        
        docx_path = os.path.join(self.temp_dir, filename)
        doc.save(docx_path)
        return docx_path
    
    def test_render_docx_to_image_basic(self):
        """Test basic .docx to image rendering."""
        docx_path = self.create_formatted_docx()
        
        # Process template
        recipient_data = {'name': 'John', 'company': 'ACME', 'email': 'john@acme.com'}
        processed_doc = self.document_processor.process_docx_template(docx_path, recipient_data)
        
        # Render to image
        image_bytes = self.document_processor.render_docx_to_image(processed_doc, 400, 300)
        
        # Verify image was generated
        assert len(image_bytes) > 0
        assert isinstance(image_bytes, bytes)
        
        # Verify it's a valid PNG
        assert image_bytes.startswith(b'\x89PNG')
    
    def test_render_docx_to_image_with_formatting(self):
        """Test that formatting is preserved in image rendering."""
        docx_path = self.create_formatted_docx()
        
        # Process template
        recipient_data = {'name': 'John Doe', 'company': 'ACME Corp', 'email': 'john@acme.com'}
        processed_doc = self.document_processor.process_docx_template(docx_path, recipient_data)
        
        # Render to image with different sizes
        small_image = self.document_processor.render_docx_to_image(processed_doc, 200, 150)
        large_image = self.document_processor.render_docx_to_image(processed_doc, 600, 400)
        
        # Verify both images were generated
        assert len(small_image) > 0
        assert len(large_image) > 0
        
        # Larger image should generally be bigger (more pixels)
        assert len(large_image) > len(small_image)
    
    def test_render_empty_document(self):
        """Test rendering an empty document."""
        doc = Document()
        doc.add_paragraph("")  # Empty paragraph
        
        image_bytes = self.document_processor.render_docx_to_image(doc, 300, 200)
        
        # Should still generate a valid image (white background)
        assert len(image_bytes) > 0
        assert image_bytes.startswith(b'\x89PNG')
    
    def test_render_long_text_wrapping(self):
        """Test that long text is properly wrapped."""
        doc = Document()
        long_text = "This is a very long line of text that should be wrapped across multiple lines when rendered to an image because it exceeds the width of the target image."
        
        para = doc.add_paragraph()
        run = para.add_run(long_text)
        run.font.size = Pt(12)
        
        image_bytes = self.document_processor.render_docx_to_image(doc, 200, 300)  # Narrow width
        
        # Should generate valid image with wrapped text
        assert len(image_bytes) > 0
        assert image_bytes.startswith(b'\x89PNG')
    
    def test_render_with_template_variables(self):
        """Test rendering with template variable substitution."""
        doc = Document()
        
        # Add paragraphs with template variables
        para1 = doc.add_paragraph()
        run1 = para1.add_run('Hello {{name}}!')
        run1.bold = True
        run1.font.size = Pt(16)
        run1.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        
        para2 = doc.add_paragraph()
        run2 = para2.add_run('Welcome to {{company}}')
        run2.italic = True
        run2.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Save document
        docx_path = os.path.join(self.temp_dir, "template.docx")
        doc.save(docx_path)
        
        # Process with recipient data
        recipient_data = {'name': 'Alice Smith', 'company': 'Tech Corp'}
        processed_doc = self.document_processor.process_docx_template(docx_path, recipient_data)
        
        # Render to image
        image_bytes = self.document_processor.render_docx_to_image(processed_doc, 400, 200)
        
        # Verify image was generated
        assert len(image_bytes) > 0
        assert image_bytes.startswith(b'\x89PNG')
        
        # Verify template variables were processed (check document content)
        processed_text = self.document_processor.convert_to_plain_text(processed_doc)
        assert 'Alice Smith' in processed_text
        assert 'Tech Corp' in processed_text
        assert '{{name}}' not in processed_text
        assert '{{company}}' not in processed_text
    
    def test_render_with_different_alignments(self):
        """Test rendering with different text alignments."""
        doc = Document()
        
        # Left aligned (default)
        left_para = doc.add_paragraph()
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_para.add_run('Left aligned text')
        
        # Center aligned
        center_para = doc.add_paragraph()
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_para.add_run('Center aligned text')
        
        # Right aligned
        right_para = doc.add_paragraph()
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_para.add_run('Right aligned text')
        
        image_bytes = self.document_processor.render_docx_to_image(doc, 400, 200)
        
        # Should generate valid image with different alignments
        assert len(image_bytes) > 0
        assert image_bytes.startswith(b'\x89PNG')
    
    def test_font_fallback(self):
        """Test that font fallback works when system fonts are not available."""
        doc = Document()
        
        para = doc.add_paragraph()
        run = para.add_run('Test text with fallback font')
        run.font.size = Pt(14)
        run.bold = True
        
        # This should work even if system fonts are not available
        image_bytes = self.document_processor.render_docx_to_image(doc, 300, 100)
        
        assert len(image_bytes) > 0
        assert image_bytes.startswith(b'\x89PNG')
    
    def test_transparent_background(self):
        """Test that formatted text rendering uses transparent background."""
        doc = Document()
        
        # Add some formatted text
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('TRANSPARENT TEST')
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Render to image
        image_bytes = self.document_processor.render_docx_to_image(doc, 200, 100)
        
        # Load the image and verify transparency
        from io import BytesIO
        image = Image.open(BytesIO(image_bytes))
        
        # Should be RGBA mode for transparency
        assert image.mode == 'RGBA'
        
        # Should have transparent pixels
        alpha_channel = image.split()[-1]
        alpha_data = list(alpha_channel.getdata())
        transparent_pixels = sum(1 for alpha in alpha_data if alpha == 0)
        
        # Most pixels should be transparent (background)
        assert transparent_pixels > len(alpha_data) * 0.5  # At least 50% transparent
        
        # Should be a valid PNG with transparency
        assert image_bytes.startswith(b'\x89PNG')
    
    def test_transparency_with_multiple_colors(self):
        """Test transparency works with multiple colored text elements."""
        doc = Document()
        
        # Add multiple colored text elements
        para1 = doc.add_paragraph()
        run1 = para1.add_run('Red Text ')
        run1.font.color.rgb = RGBColor(255, 0, 0)
        run1.font.size = Pt(14)
        
        run2 = para1.add_run('Blue Text ')
        run2.font.color.rgb = RGBColor(0, 0, 255)
        run2.font.size = Pt(14)
        
        run3 = para1.add_run('Green Text')
        run3.font.color.rgb = RGBColor(0, 255, 0)
        run3.font.size = Pt(14)
        
        # Render to image
        image_bytes = self.document_processor.render_docx_to_image(doc, 300, 50)
        
        # Load and verify
        from io import BytesIO
        image = Image.open(BytesIO(image_bytes))
        
        # Should be RGBA with transparency
        assert image.mode == 'RGBA'
        
        # Check that we have both transparent and non-transparent pixels
        alpha_channel = image.split()[-1]
        alpha_data = list(alpha_channel.getdata())
        
        transparent_pixels = sum(1 for alpha in alpha_data if alpha == 0)
        non_transparent_pixels = sum(1 for alpha in alpha_data if alpha > 0)
        
        # Should have both transparent background and visible text
        assert transparent_pixels > 0      # Has transparent background
        assert non_transparent_pixels > 0  # Has visible text (any alpha > 0)
    
    def test_docx_formatting_preservation(self):
        """Test that .docx formatting (fonts, colors, styles) is preserved in image rendering."""
        doc = Document()
        
        # Create paragraph with mixed formatting
        para = doc.add_paragraph()
        
        # Bold text
        bold_run = para.add_run('Bold ')
        bold_run.bold = True
        bold_run.font.size = Pt(14)
        bold_run.font.name = 'Arial'
        bold_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Italic text
        italic_run = para.add_run('Italic ')
        italic_run.italic = True
        italic_run.font.size = Pt(12)
        italic_run.font.name = 'Helvetica'
        italic_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        
        # Regular text
        regular_run = para.add_run('Regular')
        regular_run.font.size = Pt(10)
        regular_run.font.name = 'Monaco'
        regular_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        # Render to image
        image_bytes = self.document_processor.render_docx_to_image(doc, 200, 100)
        
        # Should render without errors and produce valid image
        assert len(image_bytes) > 0
        assert image_bytes.startswith(b'\x89PNG')
        
        # Load and verify basic properties
        from io import BytesIO
        image = Image.open(BytesIO(image_bytes))
        
        assert image.mode == 'RGBA'
        assert image.size == (200, 100)
        
        # Should have both transparent and non-transparent pixels
        alpha_channel = image.split()[-1]
        alpha_data = list(alpha_channel.getdata())
        
        transparent_pixels = sum(1 for alpha in alpha_data if alpha == 0)
        non_transparent_pixels = sum(1 for alpha in alpha_data if alpha > 0)
        
        # Should have transparent background and visible formatted text
        assert transparent_pixels > len(alpha_data) * 0.5  # Mostly transparent background
        assert non_transparent_pixels > 0  # Has visible text
    
    def test_font_family_preservation(self):
        """Test that different font families are properly loaded and used."""
        doc = Document()
        
        # Test different font families
        font_tests = [
            ('Arial', 'Arial Test'),
            ('Helvetica', 'Helvetica Test'),
            ('Monaco', 'Monaco Test'),
        ]
        
        for font_name, text in font_tests:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(12)
        
        # Should render without errors
        image_bytes = self.document_processor.render_docx_to_image(doc, 300, 150)
        
        assert len(image_bytes) > 0
        assert image_bytes.startswith(b'\x89PNG')
        
        # Load and verify
        from io import BytesIO
        image = Image.open(BytesIO(image_bytes))
        
        assert image.mode == 'RGBA'
        assert image.size == (300, 150)
    
    def test_run_structure_preservation_with_variables(self):
        """Test that run structure and formatting is preserved when processing template variables."""
        doc = Document()
        
        # Create paragraph with mixed formatting and variables
        para = doc.add_paragraph()
        
        # Regular text
        regular_run = para.add_run('Hello ')
        regular_run.font.size = Pt(14)
        regular_run.font.name = 'Arial'
        
        # Bold red text with variable
        bold_run = para.add_run('{{name}}')
        bold_run.bold = True
        bold_run.font.size = Pt(16)
        bold_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        bold_run.font.name = 'Helvetica'
        
        # Italic green text
        italic_run = para.add_run(' from ')
        italic_run.italic = True
        italic_run.font.size = Pt(12)
        italic_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        italic_run.font.name = 'Monaco'
        
        # Purple text with variable
        end_run = para.add_run('{{company}}!')
        end_run.font.size = Pt(14)
        end_run.font.color.rgb = RGBColor(128, 0, 128)  # Purple
        end_run.font.name = 'Arial'
        
        # Save to temporary file
        import tempfile
        import os
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        try:
            # Process template
            recipient_data = {'name': 'John Doe', 'company': 'ACME Corp'}
            processed_doc = self.document_processor.process_docx_template(tmp_path, recipient_data)
            
            # Verify run structure is preserved
            assert len(processed_doc.paragraphs) == 1
            para = processed_doc.paragraphs[0]
            
            # Should have 4 runs
            assert len(para.runs) == 4
            
            # Check that formatting is preserved
            runs = para.runs
            
            # Run 0: 'Hello ' - Arial, 14pt, black
            assert runs[0].text == 'Hello '
            assert runs[0].font.name == 'Arial'
            assert runs[0].font.size == Pt(14)
            assert runs[0].bold is None or runs[0].bold == False
            assert runs[0].italic is None or runs[0].italic == False
            
            # Run 1: 'John Doe' - Helvetica, 16pt, red, bold
            assert runs[1].text == 'John Doe'
            assert runs[1].font.name == 'Helvetica'
            assert runs[1].font.size == Pt(16)
            assert runs[1].bold == True
            assert str(runs[1].font.color.rgb) == 'FF0000'
            
            # Run 2: ' from ' - Monaco, 12pt, green, italic
            assert runs[2].text == ' from '
            assert runs[2].font.name == 'Monaco'
            assert runs[2].font.size == Pt(12)
            assert runs[2].italic == True
            assert str(runs[2].font.color.rgb) == '008000'
            
            # Run 3: 'ACME Corp!' - Arial, 14pt, purple
            assert runs[3].text == 'ACME Corp!'
            assert runs[3].font.name == 'Arial'
            assert runs[3].font.size == Pt(14)
            assert str(runs[3].font.color.rgb) == '800080'
            
            # Verify full text is correct
            assert para.text == 'Hello John Doe from ACME Corp!'
            
        finally:
            # Clean up
            os.unlink(tmp_path)
