"""Retrieval manager and document processing functionality."""

import os
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from urllib.parse import urlparse

from justllms.retrieval.embeddings import EmbeddingProvider, create_embedding_provider
from justllms.retrieval.models import (
    CollectionInfo,
    Document,
    ProcessingResult,
    RetrievalConfig,
    RetrievalResult,
)
from justllms.retrieval.vector_stores import VectorStore, create_vector_store


class DocumentProcessor:
    """Handles document processing and text chunking."""
    
    def __init__(self, config: RetrievalConfig):
        self.config = config
    
    def process_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process text into chunks."""
        if self.config.clean_text:
            text = self._clean_text(text)
        
        chunks = self._chunk_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = (metadata or {}).copy()
            doc_metadata.update({
                "chunk_index": i,
                "chunk_count": len(chunks),
                "chunk_size": len(chunk),
            })
            
            documents.append(Document(
                content=chunk,
                metadata=doc_metadata
            ))
        
        return documents
    
    def process_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process a file into documents."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract text based on file type
        text = self._extract_text_from_file(file_path)
        
        # Add file metadata
        file_metadata = (metadata or {}).copy()
        file_metadata.update({
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": file_path.suffix.lower(),
            "file_size": file_path.stat().st_size,
        })
        
        if self.config.extract_metadata:
            extracted_metadata = self._extract_file_metadata(file_path, text)
            file_metadata.update(extracted_metadata)
        
        return self.process_text(text, file_metadata)
    
    def process_url(self, url: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Process a URL into documents."""
        text = self._extract_text_from_url(url)
        
        url_metadata = (metadata or {}).copy()
        url_metadata.update({
            "source": url,
            "source_type": "url",
            "domain": urlparse(url).netloc,
        })
        
        return self.process_text(text, url_metadata)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        import re
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        # Remove control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        return text
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks based on strategy."""
        if self.config.splitting_strategy == "fixed":
            return self._chunk_fixed_size(text)
        elif self.config.splitting_strategy == "semantic":
            return self._chunk_semantic(text)
        else:  # recursive
            return self._chunk_recursive(text)
    
    def _chunk_fixed_size(self, text: str) -> List[str]:
        """Split text into fixed-size chunks."""
        chunks = []
        chunk_size = self.config.chunk_size
        overlap = self.config.chunk_overlap
        
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap if overlap > 0 else end
            
            if start >= len(text):
                break
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def _chunk_recursive(self, text: str) -> List[str]:
        """Recursively split text on different separators."""
        separators = ["\n\n", "\n", ". ", " ", ""]
        return self._split_text_recursive(text, separators, 0)
    
    def _split_text_recursive(self, text: str, separators: List[str], sep_index: int) -> List[str]:
        """Recursively split text using separators."""
        if len(text) <= self.config.chunk_size:
            return [text] if text.strip() else []
        
        if sep_index >= len(separators):
            # Last resort: split at chunk size
            return self._chunk_fixed_size(text)
        
        separator = separators[sep_index]
        if separator == "":
            # Character-level splitting
            return self._chunk_fixed_size(text)
        
        splits = text.split(separator)
        chunks = []
        current_chunk = ""
        
        for split in splits:
            if len(current_chunk + separator + split) <= self.config.chunk_size:
                current_chunk = current_chunk + separator + split if current_chunk else split
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                
                if len(split) > self.config.chunk_size:
                    # Split is too large, recursively split it
                    chunks.extend(self._split_text_recursive(split, separators, sep_index + 1))
                    current_chunk = ""
                else:
                    current_chunk = split
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def _chunk_semantic(self, text: str) -> List[str]:
        """Split text semantically (placeholder - would use NLP)."""
        # For now, fall back to recursive splitting
        # In a full implementation, this would use sentence embeddings
        # to find semantic boundaries
        return self._chunk_recursive(text)
    
    def _extract_text_from_file(self, file_path: Path) -> str:
        """Extract text from various file types."""
        suffix = file_path.suffix.lower()
        
        if suffix == '.txt':
            return file_path.read_text(encoding='utf-8')
        
        elif suffix == '.pdf':
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except ImportError:
                raise ImportError("PyPDF2 required for PDF processing. Install with: pip install PyPDF2")
        
        elif suffix in ['.doc', '.docx']:
            try:
                import docx
                doc = docx.Document(file_path)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                raise ImportError("python-docx required for Word processing. Install with: pip install python-docx")
        
        elif suffix in ['.html', '.htm']:
            try:
                from bs4 import BeautifulSoup
                html = file_path.read_text(encoding='utf-8')
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
            except ImportError:
                raise ImportError("beautifulsoup4 required for HTML processing. Install with: pip install beautifulsoup4")
        
        else:
            # Try to read as text
            try:
                return file_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                raise ValueError(f"Unsupported file type: {suffix}")
    
    def _extract_text_from_url(self, url: str) -> str:
        """Extract text from a URL."""
        try:
            import httpx
            from bs4 import BeautifulSoup
            
            with httpx.Client() as client:
                response = client.get(url)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.text, 'html.parser')
                return soup.get_text()
        except ImportError:
            raise ImportError("httpx and beautifulsoup4 required for URL processing")
    
    def _extract_file_metadata(self, file_path: Path, text: str) -> Dict[str, Any]:
        """Extract metadata from file content."""
        metadata = {}
        
        # Basic text statistics
        metadata["word_count"] = len(text.split())
        metadata["char_count"] = len(text)
        
        # Try to extract title (first line that looks like a title)
        lines = text.split('\n')[:10]  # Check first 10 lines
        for line in lines:
            line = line.strip()
            if line and len(line) < 200 and not line.endswith('.'):
                metadata["title"] = line
                break
        
        return metadata


class RetrievalManager:
    """Manages retrieval operations and coordinates components."""
    
    def __init__(self, config: RetrievalConfig):
        self.config = config
        
        # Initialize embedding provider
        if config.embedding_provider:
            # Use embedding_config if provided, otherwise fall back to vector_store config
            if config.embedding_config:
                embedding_config = config.embedding_config.copy()
            else:
                embedding_config = config.vector_store.copy()
                embedding_config.pop("type", None)
            
            self.embedding_provider = create_embedding_provider(
                config.embedding_provider,
                **embedding_config
            )
        else:
            self.embedding_provider = None
        
        # Initialize vector store
        store_config = config.vector_store.copy()
        store_type = store_config.pop("type")
        self.vector_store = create_vector_store(store_type, **store_config)
        
        # Initialize document processor
        self.processor = DocumentProcessor(config)
    
    def search(
        self,
        query: str,
        collection: str,
        k: Optional[int] = None,
        filters: Optional[Dict[str, Any]] = None,
        similarity_threshold: Optional[float] = None,
        **kwargs: Any
    ) -> RetrievalResult:
        """Search for relevant documents."""
        k = k or self.config.default_k
        threshold = similarity_threshold or self.config.similarity_threshold
        
        if self.embedding_provider:
            result = self.vector_store.search_with_query(
                query=query,
                embedding_provider=self.embedding_provider,
                collection=collection,
                k=k,
                filters=filters,
                **kwargs
            )
        else:
            # Fallback to direct text search if no embedding provider
            result = self.vector_store.search_text(
                query=query,
                collection=collection,
                k=k,
                filters=filters,
                **kwargs
            )
        
        # Filter by similarity threshold
        if threshold > 0:
            filtered_docs = [doc for doc in result.documents if (doc.score or 0) >= threshold]
            result.documents = filtered_docs
            result.total_results = len(filtered_docs)
        
        return result
    
    async def asearch(
        self,
        query: str,
        collection: str,
        k: Optional[int] = None,
        filters: Optional[Dict[str, Any]] = None,
        similarity_threshold: Optional[float] = None,
        **kwargs: Any
    ) -> RetrievalResult:
        """Search for relevant documents (async)."""
        k = k or self.config.default_k
        threshold = similarity_threshold or self.config.similarity_threshold
        
        start_time = time.time()
        
        # Generate embedding
        query_vector = await self.embedding_provider.aembed_text(query)
        
        # Search vector store
        documents = await self.vector_store.asearch(
            query_vector=query_vector,
            collection=collection,
            k=k,
            filters=filters,
            **kwargs
        )
        
        search_time_ms = (time.time() - start_time) * 1000
        
        # Filter by similarity threshold
        if threshold > 0:
            documents = [doc for doc in documents if (doc.score or 0) >= threshold]
        
        return RetrievalResult(
            documents=documents,
            query=query,
            total_results=len(documents),
            search_time_ms=search_time_ms
        )
    
    def add_texts(
        self,
        texts: List[Union[str, Dict[str, Any]]],
        collection: str,
        **kwargs: Any
    ) -> ProcessingResult:
        """Add text documents to collection."""
        start_time = time.time()
        
        documents = []
        for text_item in texts:
            if isinstance(text_item, str):
                docs = self.processor.process_text(text_item)
            else:
                content = text_item.get("content", "")
                metadata = text_item.get("metadata", {})
                docs = self.processor.process_text(content, metadata)
            documents.extend(docs)
        
        # Generate embeddings (if embedding provider is available)
        if self.embedding_provider:
            documents = self.embedding_provider.embed_documents(documents)
        
        # Add to vector store
        document_ids = self.vector_store.add_documents(documents, collection)
        
        processing_time_ms = (time.time() - start_time) * 1000
        
        return ProcessingResult(
            documents_processed=len(texts),
            chunks_created=len(documents),
            vectors_generated=len(documents),
            processing_time_ms=processing_time_ms
        )
    
    def ingest_documents(
        self,
        sources: List[str],
        collection: str,
        metadata: Optional[Dict[str, Any]] = None,
        **kwargs: Any
    ) -> ProcessingResult:
        """Ingest documents from files or URLs."""
        start_time = time.time()
        
        all_documents = []
        errors = []
        
        for source in sources:
            try:
                if source.startswith(('http://', 'https://')):
                    docs = self.processor.process_url(source, metadata)
                elif os.path.isdir(source):
                    docs = []
                    for root, _, files in os.walk(source):
                        for file in files:
                            file_path = os.path.join(root, file)
                            try:
                                file_docs = self.processor.process_file(file_path, metadata)
                                docs.extend(file_docs)
                            except Exception as e:
                                errors.append(f"Error processing {file_path}: {str(e)}")
                else:
                    docs = self.processor.process_file(source, metadata)
                
                all_documents.extend(docs)
                
            except Exception as e:
                errors.append(f"Error processing {source}: {str(e)}")
        
        if all_documents:
            # Generate embeddings (if embedding provider is available)
            if self.embedding_provider:
                all_documents = self.embedding_provider.embed_documents(all_documents)
            
            # Add to vector store
            document_ids = self.vector_store.add_documents(all_documents, collection)
        
        processing_time_ms = (time.time() - start_time) * 1000
        
        return ProcessingResult(
            documents_processed=len(sources),
            chunks_created=len(all_documents),
            vectors_generated=len(all_documents),
            processing_time_ms=processing_time_ms,
            errors=errors
        )
    
    def create_collection(
        self,
        name: str,
        dimension: Optional[int] = None,
        **kwargs: Any
    ) -> bool:
        """Create a new collection."""
        # Determine embedding dimension
        if dimension is None:
            if self.embedding_provider:
                sample_embedding = self.embedding_provider.embed_text("sample text")
                dimension = len(sample_embedding)
            else:
                # Default dimension for when using Pinecone's built-in embeddings
                dimension = kwargs.get("dimension", 1536)  # Common default
        
        return self.vector_store.create_collection(name, dimension, **kwargs)
    
    def list_collections(self) -> List[CollectionInfo]:
        """List all collections."""
        return self.vector_store.list_collections()
    
    def get_collection_info(self, name: str) -> Optional[CollectionInfo]:
        """Get collection information."""
        return self.vector_store.get_collection_info(name)
    
    def delete_documents(
        self,
        document_ids: List[str],
        collection: str,
        **kwargs: Any
    ) -> bool:
        """Delete documents from collection."""
        return self.vector_store.delete_documents(document_ids, collection, **kwargs)