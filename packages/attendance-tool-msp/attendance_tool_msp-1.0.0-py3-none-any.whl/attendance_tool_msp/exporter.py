import docx, docx.shared, docx.oxml, os
from datetime import datetime
from docx2pdf import convert


class Exporter:
    """
    Document export class for generating Word and PDF attendance reports.

    Design Note:
        Helper methods are instance methods because they work together
        as a cohesive workflow. Using individual helper functions as static methods
        wouldn't make sense since they depend on each other and the instance data.
        These helper methods are meant to work together internally, not to be called from
        outside, and would likely cause unexpected results if used independently.

    Attributes:
        valid_rows (list): Valid attendance records
        invalid_rows (list): Invalid records with error messages
        title (str): Document title
    """

    # Constructor with valid and invalid rows, and the document's title
    def __init__(self, valid_rows, invalid_rows, title="Attendance Report"):
        """
        Initialize the Exporter with attendance data.

        Args:
            valid_rows (list): Valid attendance records
            invalid_rows (list): Invalid records with error info
            title (str, optional): Document title. Defaults to "Attendance Report"

        Raises:
            ValueError: If valid_rows/invalid_rows are not lists or contain non-dictionary elements,
                       or if title is not a string or is empty after trimming whitespace
        """
        self.valid_rows = valid_rows
        self.invalid_rows = invalid_rows
        self.title = title

    # Getter for valid_rows
    @property
    def valid_rows(self):
        """
        Get the list of valid attendance rows.

        Returns:
            list: List of dictionaries containing valid attendance data
        """
        return self._valid_rows

    # Setter for valid_rows: Empty Lists are allowed
    @valid_rows.setter
    def valid_rows(self, valid_rows):
        """
        Set valid attendance rows with type and structure validation.

        Args:
            valid_rows (list): List of dictionaries containing valid attendance data

        Returns:
            None: This setter does not return a value

        Raises:
            ValueError: If not a list or contains non-dictionary elements
        """
        if not isinstance(valid_rows, list):
            raise ValueError("Valid rows must be a list")
        # Check if list is not empty and first item is dict
        if valid_rows and not isinstance(valid_rows[0], dict):
            raise ValueError("Valid rows data must be dictionaries")
        self._valid_rows = valid_rows

    # Getter for invalid_rows
    @property
    def invalid_rows(self):
        """
        Get the list of invalid attendance rows.

        Returns:
            list: List of dictionaries containing invalid attendance data with error messages
        """
        return self._invalid_rows

    # Setter for invalid_rows: Empty Lists are allowed
    @invalid_rows.setter
    def invalid_rows(self, invalid_rows):
        """
        Set invalid attendance rows with type and structure validation.

        Args:
            invalid_rows (list): List of dictionaries containing invalid attendance data

        Returns:
            None: This setter does not return a value

        Raises:
            ValueError: If not a list or contains non-dictionary elements
        """
        if not isinstance(invalid_rows, list):
            raise ValueError("Invalid rows must be a list")
        # Check if list is not empty and first item is dict
        if invalid_rows and not isinstance(invalid_rows[0], dict):
            raise ValueError("Invalid rows data must be dictionaries")
        self._invalid_rows = invalid_rows

    # Getter for title
    @property
    def title(self):
        """
        Get the document title.

        Returns:
            str: The document title used for headers and filenames
        """
        return self._title

    # Setter for title
    @title.setter
    def title(self, title):
        """
        Set document title with validation.

        Args:
            title (str): Document title for headers and filenames

        Returns:
            None: This setter does not return a value

        Raises:
            ValueError: If not a string or empty after trimming whitespace
        """
        if not isinstance(title, str):
            raise ValueError("Title must be a string")
        if not title.strip():
            raise ValueError("Title cannot be empty")
        self._title = title.strip()

    def export_word(self):
        """
        Generate a Word document containing attendance data.

        Returns:
            str: The file path of the generated Word document

        Raises:
            PermissionError: If document cannot be created or saved
        """
        # Initialize Document
        document = docx.Document()

        # Set page margins to 1 inch on all sides
        for section in document.sections:
            section.top_margin = docx.shared.Inches(1)
            section.bottom_margin = docx.shared.Inches(1)
            section.left_margin = docx.shared.Inches(1)
            section.right_margin = docx.shared.Inches(1)

        # Add and style document heading
        self.__add_document_heading(document)

        # Add spacing after the heading
        document.add_paragraph()

        # Create and setup the attendance table
        table = self.__create_attendance_table(document)

        # Add valid data rows
        self.__add_data_rows(table)

        # Add invalid data rows (highlighted in red)
        self.__add_data_rows(table, False)

        # Add error log section if there are invalid rows
        if self.invalid_rows:
            self.__add_error_log(document)

        # Generate unique filename with title and timestamp
        filename = self.__generate_filename()

        try:
            document.save(filename)
            return filename
        except PermissionError as error:
            # Raised possibly because file is open, and we're trying to save it
            raise PermissionError(error)

    def export_pdf(self):
        """
        Generate a PDF document by converting a Word document.
        The temporary Word document is automatically deleted after conversion.

        Returns:
            str: The file path of the generated PDF document

        Raises:
            PermissionError: If Word document creation fails due to file access issues
            Exception: If Word document creation fails or PDF conversion fails
        """
        # First create the Word document
        try:
            word_filename = self.export_word()
        except PermissionError as error:
            raise PermissionError(
                f"Failed to create Word document for PDF conversion: {error}"
            )
        except Exception as error:
            raise Exception(f"Error creating Word document for PDF conversion: {error}")

        # Convert .docx to .pdf (replace extension)
        pdf_filename = word_filename.replace(".docx", ".pdf")
        try:
            # Utilize docx2pdf package's convert method
            convert(word_filename, pdf_filename)
        except Exception as error:
            # Check if PDF was actually created despite the error
            if os.path.exists(pdf_filename):
                # This is a common issue with docx2pdf on Windows!
                # The error com_error(-2147023170, 'The remote procedure call failed.')
                # often happens when Microsoft Word is busy or has COM interface issues,
                # but the conversion actually succeeds.
                print(f"Conversion completed with minor issues: {error}")
            # Clean up & delete Word file if conversion truly failed
            else:
                try:
                    os.remove(word_filename)
                except OSError:
                    pass  # Ignore cleanup errors
                raise Exception(f"Failed to convert .docx to .pdf: {error}")

        # Clean up - delete the temporary Word document after successful conversion
        try:
            os.remove(word_filename)
            print(f"Temporary Word document '{word_filename}' deleted.")
        except OSError as error:
            print(
                f"Warning: Could not delete temporary Word file '{word_filename}': {error}"
            )

        return pdf_filename

    def __add_document_heading(self, document):
        """
        Helper method for exporting word document - Add document heading with title and branding.

        Args:
            document (docx.Document): The Word document object

        Returns:
            None: This method modifies the document in place
        """
        # Add heading
        heading = document.add_heading(
            f"{self.title}",
        )

        # Style the heading
        heading_run = heading.runs[0]  # .runs[0] = first text chunk in the heading that we can style
        heading_run.font.name = "Arial"  # Change font family to Arial
        heading_run.font.size = docx.shared.Pt(21)  # Set font size to 21 points
        heading_run.font.color.rgb = docx.shared.RGBColor(
            0, 0, 0
        )  # Set text color to black

    def __create_attendance_table(self, document):
        """
        Helper method for exporting word document - Create the main attendance data table.

        Args:
            document (docx.Document): The Word document object

        Returns:
            docx.table.Table: The customized table object
        """
        # Define table columns
        columns = ["Name", "ID", "Course Code", "Time", "Name of the Doctor"]

        # Create table: 1 header row + len(valid_rows) rows
        table = document.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"

        # Set column widths to prevent text wrapping
        column_widths = [
            docx.shared.Inches(3.5),  # Name column - wider for long names
            docx.shared.Inches(1.5),  # ID column
            docx.shared.Inches(2.0),  # Course Code column
            docx.shared.Inches(3.5),  # Time column
            docx.shared.Inches(3.0),  # Doctor/TA Name column
        ]

        for i, width in enumerate(column_widths):
            table.columns[i].width = width

        # Set table border color to blue
        self.__set_table_border_color(table)

        # --- Header Row ---
        hdr_cells = table.rows[0].cells
        # Set margins for header cells
        self.__set_cell_margins(hdr_cells)

        for i, col_name in enumerate(columns):
            paragraph = hdr_cells[i].paragraphs[0]
            run = paragraph.add_run(col_name)
            run.font.bold = True
            run.font.name = "Roboto"
            run.font.size = docx.shared.Pt(11.5)
            # paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

        return table

    def __add_data_rows(self, table, valid=True):
        """
        Helper method for exporting word document - Add data rows to the attendance table.

        Args:
            table (docx.table.Table): The table object to add rows to
            valid (bool): True for valid rows, False for invalid rows

        Returns:
            None: This method modifies the table in place
        """
        # Choose which data to process based on valid parameter - Pythonic Ternary Operator!
        for row in self.valid_rows if valid else self.invalid_rows:
            cells = table.add_row().cells
            cells[0].text = row["Full Name"]
            cells[1].text = row["University ID"]
            cells[2].text = row["Course Code"]
            cells[3].text = row["Course Time"]
            cells[4].text = row["Doctor/TA Name"]

            # Set margins for data cells
            self.__set_cell_margins(cells)

            # Format row text
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = docx.shared.Pt(11.5)
                        run.font.name = "Roboto"
                        # Make invalid rows red
                        if not valid:
                            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

    def __set_cell_margins(self, cells):
        """
        Helper method for exporting word document - Set cell margins using XML manipulation.

        Args:
            cells (list): List of table cells to apply margins to

        Returns:
            None: This method modifies the cells in place
        """

        # Specific Margin Sizes
        top_margin = 150
        right_margin = 360
        bottom_margin = 360
        left_margin = 150

        for cell in cells:
            cell_element = cell._element
            cell_properties = cell_element.get_or_add_tcPr()
            margins = docx.oxml.parse_xml(f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                        f'<w:top w:w="{top_margin}" w:type="dxa"/>'
                                        f'<w:left w:w="{left_margin}" w:type="dxa"/>'
                                        f'<w:bottom w:w="{bottom_margin}" w:type="dxa"/>'
                                        f'<w:right w:w="{right_margin}" w:type="dxa"/>'
                                        f'</w:tcMar>')
            cell_properties.append(margins)

    def __set_table_border_color(self, table, color="0000FF"):
        """
        Helper method for exporting word document - Set table border color using XML manipulation.

        Args:
            table (docx.table.Table): The table object to apply border color to
            color (str): Hex color code (default "0000FF" for blue)

        Returns:
            None: This method modifies the table in place
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = docx.oxml.parse_xml(f'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                       f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'</w:tblBorders>')
        tblPr.append(tblBorders)

    def __add_error_log(self, document):
        """
        Helper method for exporting word document - Add validation error log section to the document.

        Args:
            document (docx.Document): The Word document to add the error log to

        Returns:
            None: This method modifies the document in place
        """
        # Add some space before the log
        document.add_paragraph()

        # Add log heading
        log_heading = document.add_paragraph()
        log_heading_run = log_heading.add_run("Validation Issues Log:")
        log_heading_run.font.bold = True
        log_heading_run.font.size = docx.shared.Pt(14)
        log_heading_run.font.name = "Roboto"
        log_heading_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red color

        # Add each error as a separate paragraph
        for i, row in enumerate(self.invalid_rows, 1):
            if "error" in row and row["error"]:
                error_paragraph = document.add_paragraph()

                # Add error number (invalid_row index + 1)
                error_run = error_paragraph.add_run(f"{i}. ")
                error_run.font.bold = True
                error_run.font.name = "Roboto"
                error_run.font.size = docx.shared.Pt(11)

                # Add student identifier (name or ID)
                student_name = row["Full Name"]
                if not student_name.strip():
                    # Default value 'Unknown'
                    if row["University ID"]:
                        student_name = f"Student with ID: {row["University ID"]}"
                    else:
                        student_name = f"Unknown Student"

                # Student -
                name_run = error_paragraph.add_run(f"{student_name} - ")
                name_run.font.bold = True
                name_run.font.name = "Roboto"
                name_run.font.size = docx.shared.Pt(11)

                # Add error message
                error_msg_run = error_paragraph.add_run(row["error"])
                error_msg_run.font.name = "Roboto"
                error_msg_run.font.size = docx.shared.Pt(11)
                error_msg_run.font.color.rgb = docx.shared.RGBColor(
                    128, 128, 128
                )  # Gray color

    def __generate_filename(self):
        """
        Helper method for exporting word document - Generate a unique filename based on title and current timestamp.

        Example:
            {title}_{date}_{hr&min&sec}.docx

        Returns:
            str: A unique filename for the Word document
        """
        # Clean title - replace spaces with underscores
        clean_title = self.title.replace(" ", "_")

        # Remove any remaining problematic characters
        safe_characters = []
        for char in clean_title:
            if char.isalnum() or char == "_":
                safe_characters.append(char)
            else:
                safe_characters.append("_")

        clean_title = "".join(safe_characters).strip("_")

        # Simple timestamp
        now = datetime.now()
        timestamp = f"{now.year}{now.month:02d}{now.day:02d}_{now.hour:02d}{now.minute:02d}{now.second:02d}"

        # Create filename
        filename = f"{clean_title}_{timestamp}.docx"

        # Use when testing, and uncomment above.
        # filename = f"demo.docx"

        return filename
