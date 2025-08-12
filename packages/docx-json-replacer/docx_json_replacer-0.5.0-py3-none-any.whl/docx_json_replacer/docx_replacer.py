"""
DocxReplacer with table support for docx-json-replacer
"""
import json
from typing import Dict, Any
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from .utility.html_parse import clean_html_content
    from .table_handler import TableHandler
except ImportError:
    from utility.html_parse import clean_html_content
    from table_handler import TableHandler


class DocxReplacer:
    """Replace placeholders in DOCX files with JSON data, including table support"""
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.table_handler = TableHandler()
        self.table_placeholders = {}
    
    def replace_from_json(self, json_data: Dict[str, Any]) -> None:
        """Replace placeholders and create tables from JSON data"""
        
        # First pass: identify table placeholders
        for paragraph in self.doc.paragraphs:
            original_text = paragraph.text
            for key, value in json_data.items():
                placeholder = "{{" + key + "}}"
                placeholder_spaced = "{{ " + key + " }}"
                
                if placeholder in paragraph.text or placeholder_spaced in paragraph.text:
                    # Check if this is table data
                    if self._is_table_data(value):
                        # Store for table insertion and clear the paragraph
                        self.table_placeholders[paragraph] = (key, value)
                        # Clear only the placeholder, not the entire paragraph
                        paragraph.text = paragraph.text.replace(placeholder, "")
                        paragraph.text = paragraph.text.replace(placeholder_spaced, "")
                    else:
                        # Regular text replacement
                        cleaned_value = clean_html_content(value)
                        paragraph.text = paragraph.text.replace(placeholder, str(cleaned_value))
                        paragraph.text = paragraph.text.replace(placeholder_spaced, str(cleaned_value))
        
        # Second pass: insert tables where placeholders were found
        for paragraph, (key, value) in self.table_placeholders.items():
            self._insert_table_after_paragraph(paragraph, value)
    
    def _is_table_data(self, value: Any) -> bool:
        """Check if value is table data"""
        return (self.table_handler.is_table_data(value) or 
                (isinstance(value, str) and '<table' in value.lower()))
    
    def _insert_table_after_paragraph(self, paragraph, value: Any) -> None:
        """Insert a formatted table after a paragraph"""
        
        # Process table data
        table_data = self.table_handler.process_table_data(value)
        
        if not table_data.get('rows'):
            return
        
        # Get parent element and index
        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)
        
        # Determine table dimensions
        rows = table_data['rows']
        num_rows = len(rows)
        num_cols = len(rows[0]['cells']) if rows and 'cells' in rows[0] else 0
        
        if num_rows == 0 or num_cols == 0:
            return
        
        # Create new table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Fill table with data and apply styling
        for row_idx, row_data in enumerate(rows):
            cells = row_data.get('cells', [])
            style = row_data.get('style', {})
            
            for col_idx, cell_text in enumerate(cells):
                if col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = str(cell_text)
                    
                    # Apply styling
                    if style:
                        self._apply_cell_style(cell, style)
        
        # Move table to correct position
        parent.insert(index + 1, table._element)
    
    def _apply_cell_style(self, cell, style: Dict[str, Any]) -> None:
        """Apply styling to a table cell"""
        
        # Apply background color
        if 'bg' in style and style['bg']:
            self._set_cell_background(cell, style['bg'])
        
        # Apply text formatting
        if cell.paragraphs:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if style.get('bold'):
                        run.bold = True
                    if style.get('italic'):
                        run.italic = True
                    
                    # Apply text color
                    if 'color' in style and style['color']:
                        color_hex = style['color'].replace('#', '')
                        if len(color_hex) == 6:
                            r = int(color_hex[0:2], 16)
                            g = int(color_hex[2:4], 16)
                            b = int(color_hex[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
    
    def _set_cell_background(self, cell, color: str) -> None:
        """Set background color for a table cell"""
        color = color.replace('#', '')
        
        # Get or create cell properties
        tc_pr = cell._tc.get_or_add_tcPr()
        
        # Create shading element
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        
        # Remove existing shading
        existing_shd = tc_pr.find(qn('w:shd'))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)
        
        # Add new shading
        tc_pr.append(shd)
    
    def save(self, output_path: str) -> None:
        """Save the modified document"""
        self.doc.save(output_path)
    
    def replace_from_json_file(self, json_path: str) -> None:
        """Load JSON from file and replace"""
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        self.replace_from_json(json_data)


def replace_docx_template(docx_path: str, json_data: Dict[str, Any], output_path: str) -> None:
    """Utility function to replace template and save in one step"""
    replacer = DocxReplacer(docx_path)
    replacer.replace_from_json(json_data)
    replacer.save(output_path)